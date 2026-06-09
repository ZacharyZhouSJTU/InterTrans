"""文档导入模块 - 支持 .docx / .txt 文件解析与句子切割 + 中英文过滤

本版改进：
- .docx 不只读取普通段落，也会读取表格、页眉页脚、文本框等 Word XML 中的文本；
- .txt 自动兼容 UTF-8 / UTF-8-SIG / GBK / GB18030；
- 上传文件读取前会 seek(0)，避免 Streamlit UploadedFile 指针被提前消费导致空读；
- 解析结果为空时返回 []，由 app.py 给出友好提示。
"""

from __future__ import annotations

import io
import re
import zipfile
import xml.etree.ElementTree as ET
from typing import Iterable

try:
    from docx import Document
except Exception:  # pragma: no cover
    Document = None


_SENTENCE_SPLITTER = re.compile(r"(?<=[。！？；.!?;])\s*(?=\S)")
_CHINESE_PATTERN = re.compile(r"[\u4e00-\u9fff\u3400-\u4dbf\uf900-\ufaff]")
_W_NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"


def parse_uploaded_file(uploaded_file) -> list[dict]:
    """
    解析上传的文件，按句子切分，生成 Sentence ID。
    """
    filename = getattr(uploaded_file, "name", "").lower()

    try:
        uploaded_file.seek(0)
    except Exception:
        pass

    if filename.endswith(".txt"):
        raw = uploaded_file.read()
        content = raw if isinstance(raw, str) else _decode_text_bytes(raw)
        raw_paragraphs = _split_paragraphs(content)
    elif filename.endswith(".docx"):
        content_bytes = uploaded_file.read()
        if not content_bytes:
            return []
        raw_paragraphs = _parse_docx(content_bytes)
    else:
        raise ValueError(f"不支持的文件格式：{getattr(uploaded_file, 'name', filename)}，请上传 .docx 或 .txt 文件")

    sentences = _split_sentences(raw_paragraphs)

    results = []
    for idx, sent in enumerate((s.strip() for s in sentences if s and s.strip()), start=1):
        results.append({
            "sentence_id": f"S{idx}",
            "source_text": sent,
        })
    return results


def _decode_text_bytes(raw: bytes) -> str:
    if raw is None:
        return ""
    for encoding in ("utf-8-sig", "utf-8", "gb18030", "gbk"):
        try:
            return raw.decode(encoding)
        except UnicodeDecodeError:
            continue
    return raw.decode("utf-8", errors="replace")


def _split_paragraphs(content: str) -> list[str]:
    if not content:
        return []
    content = content.replace("\r\n", "\n").replace("\r", "\n")
    rough_parts = re.split(r"\n\s*\n+", content)
    paragraphs: list[str] = []
    for part in rough_parts:
        for line in part.split("\n"):
            line = _normalize_text(line)
            if line:
                paragraphs.append(line)
    return paragraphs


def _parse_docx(content_bytes: bytes) -> list[str]:
    paragraphs = _parse_docx_word_xml(content_bytes)
    if not paragraphs:
        paragraphs = _parse_docx_with_python_docx(content_bytes)
    return _dedupe_preserve_order(paragraphs)


def _parse_docx_word_xml(content_bytes: bytes) -> list[str]:
    """
    直接解析 Word XML，覆盖普通段落、表格、文本框、页眉页脚、脚注尾注等。
    """
    try:
        with zipfile.ZipFile(io.BytesIO(content_bytes)) as zf:
            names = zf.namelist()
            preferred = ["word/document.xml"]
            extras = sorted(
                n for n in names
                if n.startswith("word/")
                and n.endswith(".xml")
                and (
                    "/header" in n
                    or "/footer" in n
                    or n.endswith("footnotes.xml")
                    or n.endswith("endnotes.xml")
                    or n.endswith("comments.xml")
                )
            )
            xml_names = [n for n in preferred + extras if n in names]

            paragraphs: list[str] = []
            for name in xml_names:
                try:
                    root = ET.fromstring(zf.read(name))
                except Exception:
                    continue
                for para in root.iter(_W_NS + "p"):
                    text = _extract_text_from_w_p(para)
                    text = _normalize_text(text)
                    if text:
                        paragraphs.append(text)

            return _dedupe_preserve_order(paragraphs)
    except Exception:
        return []


def _extract_text_from_w_p(para: ET.Element) -> str:
    parts: list[str] = []
    for node in para.iter():
        tag = node.tag
        if tag == _W_NS + "t":
            parts.append(node.text or "")
        elif tag == _W_NS + "tab":
            parts.append("\t")
        elif tag in {_W_NS + "br", _W_NS + "cr"}:
            parts.append("\n")
    return "".join(parts)


def _parse_docx_with_python_docx(content_bytes: bytes) -> list[str]:
    if Document is None:
        return []
    try:
        doc = Document(io.BytesIO(content_bytes))
    except Exception:
        return []

    paragraphs: list[str] = []

    for para in doc.paragraphs:
        text = _normalize_text(para.text)
        if text:
            paragraphs.append(text)

    for table in doc.tables:
        paragraphs.extend(_extract_table_text(table))

    for section in doc.sections:
        for container in (section.header, section.footer):
            for para in container.paragraphs:
                text = _normalize_text(para.text)
                if text:
                    paragraphs.append(text)
            for table in container.tables:
                paragraphs.extend(_extract_table_text(table))

    return _dedupe_preserve_order(paragraphs)


def _extract_table_text(table) -> list[str]:
    paragraphs: list[str] = []
    for row in table.rows:
        cell_texts: list[str] = []
        for cell in row.cells:
            inner_parts: list[str] = []
            for para in cell.paragraphs:
                text = _normalize_text(para.text)
                if text:
                    inner_parts.append(text)
            for nested in cell.tables:
                inner_parts.extend(_extract_table_text(nested))
            cell_text = _normalize_text(" \t ".join(inner_parts))
            if cell_text:
                cell_texts.append(cell_text)
        row_text = _normalize_text(" \t ".join(cell_texts))
        if row_text:
            paragraphs.append(row_text)
    return paragraphs


def _split_sentences(paragraphs: Iterable[str]) -> list[str]:
    sentences: list[str] = []
    for para in paragraphs or []:
        para = _normalize_text(str(para))
        if not para:
            continue
        tab_parts = [p for p in re.split(r"\t+", para) if p and p.strip()]
        for tab_part in tab_parts or [para]:
            parts = _SENTENCE_SPLITTER.split(tab_part)
            for part in parts:
                s = _normalize_text(part)
                if s:
                    sentences.append(s)
    return sentences


def _normalize_text(text: str) -> str:
    if text is None:
        return ""
    text = str(text)
    text = text.replace("\u00a0", " ").replace("\u3000", " ")
    text = re.sub(r"[ \f\v]+", " ", text)
    text = re.sub(r"\n\s*\n+", "\n", text)
    return text.strip()


def _dedupe_preserve_order(items: Iterable[str]) -> list[str]:
    seen: set[str] = set()
    result: list[str] = []
    for item in items or []:
        text = _normalize_text(item)
        if not text:
            continue
        key = re.sub(r"\s+", " ", text).strip()
        if key in seen:
            continue
        seen.add(key)
        result.append(text)
    return result


# ═══════════════════════════════════════════════════════════════
#  中英文检测与过滤
# ═══════════════════════════════════════════════════════════════

def has_chinese(text: str) -> bool:
    return bool(_CHINESE_PATTERN.search(text or ""))


def filter_chinese_only(segments: list[dict]) -> list[dict]:
    return [seg for seg in segments if has_chinese(seg.get("source_text", ""))]


def filter_english_only(segments: list[dict]) -> list[dict]:
    return [seg for seg in segments if not has_chinese(seg.get("source_text", ""))]


# ═══════════════════════════════════════════════════════════════
#  Excel 导出（自动换行 + 自适应列宽）
# ═══════════════════════════════════════════════════════════════

def export_to_xlsx(segments: list[dict], col_label: str = "Source Text") -> bytes:
    """
    导出 Excel。openpyxl 改为懒加载，避免部署环境缺少 openpyxl 时导致整个网页启动失败。
    """
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Alignment
        from openpyxl.utils import get_column_letter
    except ModuleNotFoundError as exc:
        raise RuntimeError(
            "缺少 openpyxl。请在 requirements.txt 中确认包含 openpyxl>=3.0.0，"
            "然后重新部署；本地可运行：pip install openpyxl"
        ) from exc

    wb = Workbook()
    ws = wb.active
    ws.title = "Sentences"

    headers = ["Sentence ID", col_label]
    ws.append(headers)

    for cell in ws[1]:
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)

    for seg in segments:
        ws.append([seg.get("sentence_id", ""), seg.get("source_text", "")])

    ws.column_dimensions["A"].width = 14

    max_len = 0
    for seg in segments:
        for line in str(seg.get("source_text", "")).split("\n"):
            line_len = sum(2 if ord(c) > 127 else 1 for c in line)
            max_len = max(max_len, line_len)

    col_width = max(min(max_len + 4, 80), 40)
    ws.column_dimensions[get_column_letter(2)].width = col_width

    for row in ws.iter_rows(min_row=2, max_row=ws.max_row, min_col=2, max_col=2):
        for cell in row:
            cell.alignment = Alignment(vertical="top", wrap_text=True)

    ws.freeze_panes = "A2"

    output = io.BytesIO()
    wb.save(output)
    output.seek(0)
    return output.getvalue()
